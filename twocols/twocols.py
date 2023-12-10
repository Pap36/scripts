from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_column_paragraph(table, text, translation):
    # Add a new row to the table
    row = table.add_row()

    # Add English text to the left column
    cell_left = row.cells[0]
    cell_left.text = text
    paragraph_left = cell_left.paragraphs[0]
    paragraph_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add translation to the right column
    cell_right = row.cells[1]
    cell_right.text = translation
    paragraph_right = cell_right.paragraphs[0]
    paragraph_right.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# English text
with open('eng.txt', 'r') as file:
    english_text = file.read().splitlines()

# Translations
with open('rom.txt', 'r') as file:
    translations = file.read().splitlines()

# Create a new Word document
doc = Document()

# Add a table with two columns
table = doc.add_table(rows=1, cols=2)
table.autofit = True  # Disable autofit to customize column width

# Set column widths (adjust as needed)
# column_width_left = Cm(4)
# column_width_right = Cm(4)
# table.columns[0].width = column_width_left
# table.columns[1].width = column_width_right

# Add a title to the table
title_row = table.rows[0]
# title_row.cells[0].text = "English"
# title_row.cells[1].text = "Translation"

# Add content to the table
for text, translation in zip(english_text, translations):
    add_column_paragraph(table, text, translation)

# Save the document
doc.save('translation_document.docx')